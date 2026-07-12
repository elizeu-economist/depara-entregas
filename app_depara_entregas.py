import streamlit as st
import pandas as pd
import numpy as np
import unicodedata
import difflib
import re
import datetime
import io
import os

from dotenv import load_dotenv
from databricks import sql as databricks_sql
from openpyxl.styles import PatternFill, Font as XLFont
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# =======================================================
# ==== Credenciais via .env ====
# =======================================================
load_dotenv()
HOST      = "dbc-9c96b0a5-aa83.cloud.databricks.com"
HTTP_PATH = "/sql/1.0/warehouses/66eb77f90d9b7746"
TOKEN     = os.getenv("DATABRICKS_TOKEN")

# =======================================================
# ==== Configuração da Página ====
# =======================================================
st.set_page_config(
    page_title="Comparador IMR vs Databricks",
    page_icon="📦",
    layout="wide"
)

# =======================================================
# ==== Login ====
# =======================================================
def verificar_login():
    if st.session_state.get("autenticado", False):
        return True

    st.title("🔒 Acesso restrito")
    st.markdown("Informe suas credenciais para acessar o comparador.")

    with st.form("form_login"):
        usuario_input = st.text_input("Usuário")
        senha_input   = st.text_input("Senha", type="password")
        entrar = st.form_submit_button("Entrar", type="primary")

    if entrar:
        usuario_ok = usuario_input == st.secrets["credenciais"]["usuario"]
        senha_ok   = senha_input == st.secrets["credenciais"]["senha"]
        if usuario_ok and senha_ok:
            st.session_state.autenticado = True
            st.rerun()
        else:
            st.error("Usuário ou senha incorretos.")

    return False


if not verificar_login():
    st.stop()

with st.sidebar:
    if st.button("🚪 Sair"):
        st.session_state.autenticado = False
        st.rerun()

st.title("📦 Comparador IMR vs Databricks")
st.markdown("Selecione o tipo de relatório, o período, faça o upload do IMR e gere o relatório de cruzamento (De-Para).")

# =======================================================
# ==== Constantes gerais ====
# =======================================================
NOMES_MESES = ["Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
               "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]

COR_HEADER_AZUL = (0, 32, 96)
COR_DIVERGENCIA = (255, 199, 206)
COR_OK          = (198, 239, 206)
COR_AVISO       = (255, 235, 156)
COR_CINZA_CLARO = (242, 242, 242)

# Dicionário de conversão Ano/Mês -> imr_referencia (usado apenas em Comprovantes).
# Quando surgir um novo mês, basta adicionar a linha aqui.
DICIONARIO_IMR_COMPROVANTES = {
    "2025-01": "7",
    "2025-02": "8",
    "2025-03": "18",
    "2025-04": "21",
    "2025-05": "24",
    "2025-06": "25",
    "2025-07": "26",
    "2025-08": "30",
    "2025-09": "31",
    "2025-10": "32",
    "2025-11": "33",
    "2025-12": "34",
    "2026-01": "36",
    "2026-02": "37",
}

# =======================================================
# ==== Configuração por tipo de relatório ====
# =======================================================
CONFIG = {
    "Entregas": {
        "emoji": "🚚",
        "aba_bruta": "Entregas",
        "exemplo_arquivo": "IMR_AGO_entregas.xlsx",
        "prefixo_saida": "Depara_entregas",
        "colunas_esperadas": [
            "numero_pedido", "data_de_corte", "data_prazo_entrega",
            "data_de_entrega", "armazem", "tipo_do_pedido",
            "destinatario", "programa_saude", "cidade_destinatario",
            "modal", "agendamento", "qualidade_de_entrega"
        ],
        "date_cols": ["data_de_corte", "data_prazo_entrega", "data_de_entrega"],
        "cols_ignorar_hora": ["data_prazo_entrega", "data_de_entrega"],
        "de_para_forcado": {},
        "key_clean_forte": False,
    },
    "Armazém - Pedidos": {
        "emoji": "🏭",
        "aba_bruta": "Armazém_Pedidos",
        "exemplo_arquivo": "IMR_AGO_armazempedidos.xlsx",
        "prefixo_saida": "Depara_armazempedidos",
        "colunas_esperadas": [
            "numero_pedido", "data_de_corte", "data_prazo_producao",
            "data_de_producao", "armazem", "tipo_do_pedido",
            "destinatario", "programa_saude", "cidade",
            "modal", "agendamento", "qualidade_de_producao"
        ],
        "date_cols": ["data_de_corte", "data_prazo_producao", "data_de_producao"],
        "cols_ignorar_hora": [],
        "de_para_forcado": {},
        "key_clean_forte": True,
    },
    "Comprovantes": {
        "emoji": "📑",
        "aba_bruta": "Comprovantes",
        "exemplo_arquivo": "IMR_AGO_comprovantes.xlsx",
        "prefixo_saida": "Depara_comprovantes",
        "colunas_esperadas": [
            "numero_pedido", "data_corte", "data_de_entrega",
            "data_prazo_comprovante", "data_comprovante", "armazem",
            "tipo_pedido", "destinatario", "programa_saude",
            "cidade_destinatario", "modal"
        ],
        "date_cols": ["data_corte", "data_de_entrega", "data_prazo_comprovante", "data_comprovante"],
        "cols_ignorar_hora": [],
        "de_para_forcado": {"cidade": "cidade_destinatario"},
        "key_clean_forte": False,
    },
    "Armazém - Quebra Depósito": {
        "emoji": "📦",
        "aba_bruta": "Armazém_Quebra_Depósito",
        "exemplo_arquivo": "IMR_AGO_quebradeposito.xlsx",
        "prefixo_saida": "Depara_quebradeposito",
    },
}

# =======================================================
# ==== Funções Utilitárias ====
# =======================================================
def limpar_nome_coluna(col):
    if not isinstance(col, str):
        return str(col)
    c = unicodedata.normalize('NFKD', col).encode('ASCII', 'ignore').decode('utf-8')
    c = re.sub(r'(?i)\bn[oº.]?\b', 'numero', c)
    c = re.sub(r'(?i)N[oº.]', 'numero', c)
    c = re.sub(r'[^a-zA-Z0-9]+', '_', c)
    return c.strip('_').lower()


def parse_dt_robust(series):
    s = pd.Series(series).copy()
    if pd.api.types.is_datetime64_any_dtype(s):
        return s.dt.tz_localize(None)
    s = s.astype(str).str.replace(r'[\r\n]', ' ', regex=True).str.strip()
    s = s.replace(['nan', 'NaT', 'None', ''], np.nan)
    is_num = s.str.isnumeric() | (
        s.str.replace('.', '', 1).str.isnumeric() & ~s.str.contains('[-/:]', na=False)
    )
    nums = pd.to_numeric(s[is_num], errors='coerce')
    s_num = pd.to_datetime(nums, unit='D', origin='1899-12-30', errors='coerce')
    s_txt = pd.to_datetime(s[~is_num], format='mixed', dayfirst=True, errors='coerce')
    return pd.concat([s_num, s_txt]).reindex(s.index)


def normalize_text(series):
    s = series.astype(str).replace(['nan', 'None', '<NA>', 'NaT'], '')
    s = s.str.replace(r'[\r\n]', '', regex=True).str.upper().str.strip()
    return s.apply(
        lambda x: unicodedata.normalize('NFKD', str(x)).encode('ASCII', 'ignore').decode('utf-8') if x else ''
    )


def compare_one_col(joined, col, key, date_cols, cols_ignorar_hora):
    nm_imr, nm_db = f"{col}_imr", f"{col}_db"
    val_imr = joined[nm_imr]
    val_db  = joined[nm_db] if nm_db in joined.columns else pd.Series([np.nan] * len(joined))

    if col in date_cols:
        match_final = (val_imr.isna() & val_db.isna()) | (
            val_imr.notna() & val_db.notna() & (val_imr == val_db)
        )
        fmt = "%d/%m/%Y" if col in cols_ignorar_hora else "%d/%m/%Y %H:%M"
        val_imr_str = val_imr.dt.strftime(fmt)
        val_db_str  = val_db.dt.strftime(fmt)
    else:
        t_imr, t_db = normalize_text(val_imr), normalize_text(val_db)
        match_final = (t_imr == t_db)
        val_imr_str = val_imr.astype(str)
        val_db_str  = val_db.astype(str)

    return pd.DataFrame({
        'Nº Pedido': joined[key],
        'coluna': col,
        'valor_imr': val_imr_str,
        'valor_db': val_db_str,
        'match_final': match_final
    })


def format_out(df, cols_ignorar_hora):
    df_copy = df.copy()
    for col in df_copy.select_dtypes(include=['datetime64', 'datetimetz']).columns:
        fmt = "%d/%m/%Y" if col in cols_ignorar_hora else "%d/%m/%Y %H:%M:%S"
        df_copy[col] = df_copy[col].dt.strftime(fmt)
    return df_copy


def ordenar_por_pedido(df, col_name):
    if df.empty or col_name not in df.columns:
        return df
    df_copy = df.copy()
    df_copy['ordem_tmp'] = pd.to_numeric(df_copy[col_name], errors='coerce')
    df_copy = df_copy.sort_values(['ordem_tmp', col_name], ascending=[True, True])
    return df_copy.drop(columns=['ordem_tmp'])


def ordenar_por_chaves(df, chaves):
    if df.empty:
        return df
    colunas_presentes = [k for k in chaves if k in df.columns]
    if not colunas_presentes:
        return df
    df_copy = df.copy()
    primeira_chave = colunas_presentes[0]
    df_copy['ordem_tmp'] = pd.to_numeric(df_copy[primeira_chave], errors='coerce')
    df_copy = df_copy.sort_values(
        ['ordem_tmp'] + colunas_presentes,
        ascending=[True] + [True] * len(colunas_presentes)
    )
    return df_copy.drop(columns=['ordem_tmp'])


def limpar_chave(series, forte=False):
    s = series.astype(str).str.strip()
    if forte:
        s = s.str.replace(r'\.0$', '', regex=True).str.upper().replace('NAN', '')
    return s


# =======================================================
# ==== Extração Databricks ====
# =======================================================
@st.cache_data(show_spinner=False)
def extrair_databricks(query):
    conexao = databricks_sql.connect(
        server_hostname=HOST,
        http_path=HTTP_PATH,
        access_token=TOKEN
    )
    df = pd.read_sql(query, conexao)
    conexao.close()
    df.columns = df.columns.str.lower()
    return df


def montar_query(tipo, ano, mes):
    """Retorna a query correta para cada tipo de relatório."""
    if tipo == "Entregas":
        return f"""
            SELECT *
            FROM vtclog.gold.cte_imr
            WHERE ano_imr_referencia = {ano}
              AND mes_imr_referencia = {mes}
        """
    if tipo == "Armazém - Pedidos":
        return f"""
            SELECT *
            FROM vtclog.gold.armazem_pedidos
            WHERE ano = {ano} AND mes = {mes}
        """
    if tipo == "Comprovantes":
        chave = f"{ano}-{mes:02d}"
        imr_ref = DICIONARIO_IMR_COMPROVANTES[chave]
        return f"""
            SELECT *
            FROM vtclog.gold.comprovantes_imr
            WHERE imr_referencia = '{imr_ref}'
        """
    if tipo == "Armazém - Quebra Depósito":
        return f"""
            SELECT *
            FROM vtclog.gold.armazem_quebra_deposito
            WHERE mes_referencia = '{ano}-{mes:02d}'
        """
    raise ValueError(f"Tipo de relatório desconhecido: {tipo}")


# =======================================================
# ==== Processamento: relatórios por Nº Pedido ====
# ==== (Entregas, Armazém - Pedidos, Comprovantes) ====
# =======================================================
def processar_pedidos(df_base, df_ref, cfg):
    colunas_esperadas = cfg["colunas_esperadas"]
    date_cols         = cfg["date_cols"]
    cols_ignorar_hora = cfg["cols_ignorar_hora"]
    de_para_forcado   = cfg["de_para_forcado"]

    qtd_linhas_imr = len(df_base)
    qtd_linhas_db  = len(df_ref)

    # Padronização automática de cabeçalhos (fuzzy)
    mapa_colunas_imr = {}
    for col_original in df_base.columns:
        col_limpa = limpar_nome_coluna(col_original)
        if col_limpa in de_para_forcado:
            mapa_colunas_imr[col_original] = de_para_forcado[col_limpa]
        else:
            matches = difflib.get_close_matches(col_limpa, colunas_esperadas, n=1, cutoff=0.6)
            if matches:
                mapa_colunas_imr[col_original] = matches[0]
    df_base = df_base.rename(columns=mapa_colunas_imr)

    key = "numero_pedido"
    if key not in df_base.columns or key not in df_ref.columns:
        raise ValueError(f"Coluna chave '{key}' não encontrada. Verifique os cabeçalhos.")

    df_base[key] = limpar_chave(df_base[key], forte=cfg["key_clean_forte"])
    df_ref[key]  = limpar_chave(df_ref[key],  forte=cfg["key_clean_forte"])

    for col in date_cols:
        if col in df_base.columns:
            df_base[col] = parse_dt_robust(df_base[col])
            df_base[col] = df_base[col].dt.floor('D') if col in cols_ignorar_hora else df_base[col].dt.floor('min')
        if col in df_ref.columns:
            df_ref[col] = parse_dt_robust(df_ref[col])
            df_ref[col] = df_ref[col].dt.floor('D') if col in cols_ignorar_hora else df_ref[col].dt.floor('min')

    joined = pd.merge(df_base, df_ref, on=key, how='inner', suffixes=('_imr', '_db'))
    cols_to_compare = [c for c in colunas_esperadas if c != key and c in df_base.columns]

    if len(joined) > 0:
        comparacao = pd.concat(
            [compare_one_col(joined, col, key, date_cols, cols_ignorar_hora) for col in cols_to_compare],
            ignore_index=True
        )
    else:
        comparacao = pd.DataFrame(columns=['Nº Pedido', 'coluna', 'valor_imr', 'valor_db', 'match_final'])

    stats_colunas = comparacao.groupby('coluna').agg(
        total_comparacoes=('match_final', 'count'),
        verdadeiro_n=('match_final', lambda x: x.sum()),
        falso_n=('match_final', lambda x: (~x).sum())
    ).reset_index()
    stats_colunas['pct_v_num'] = (stats_colunas['verdadeiro_n'] / stats_colunas['total_comparacoes']) * 100
    stats_colunas['pct_f_num'] = (stats_colunas['falso_n'] / stats_colunas['total_comparacoes']) * 100
    stats_colunas['pct_verdadeiro'] = stats_colunas['pct_v_num'].apply(lambda x: f"{x:,.2f}%".replace('.', ','))
    stats_colunas['pct_falso']      = stats_colunas['pct_f_num'].apply(lambda x: f"{x:,.2f}%".replace('.', ','))
    stats_colunas = stats_colunas[['coluna', 'total_comparacoes', 'verdadeiro_n', 'falso_n', 'pct_verdadeiro', 'pct_falso']]

    linhas_resumo = pd.DataFrame({
        'coluna': ["Qtd Linhas IMR Total", "Qtd Linhas DB Total"],
        'total_comparacoes': [qtd_linhas_imr, qtd_linhas_db],
        'verdadeiro_n': [np.nan, np.nan], 'falso_n': [np.nan, np.nan],
        'pct_verdadeiro': [np.nan, np.nan], 'pct_falso': [np.nan, np.nan]
    })
    stats_final = pd.concat([stats_colunas, linhas_resumo], ignore_index=True)

    # Painel Match — lógica padrão (Entregas): ignora apenas 'armazem'
    if not comparacao.empty:
        painel_match = comparacao.pivot_table(
            index='Nº Pedido', columns='coluna', values='match_final',
            aggfunc='first', fill_value=False
        ).reset_index()
        colunas_teste = [c for c in painel_match.columns if c not in ['Nº Pedido', 'armazem']]
        mascara_divergencias = (painel_match[colunas_teste] == False).any(axis=1)
        painel_match = painel_match[mascara_divergencias]
        ordem_desejada = ['Nº Pedido'] + [c for c in df_base.columns if c in painel_match.columns and c != key]
        painel_match = painel_match[[c for c in ordem_desejada if c in painel_match.columns]]
    else:
        painel_match = pd.DataFrame()

    somente_no_db  = format_out(df_ref[~df_ref[key].isin(df_base[key])], cols_ignorar_hora)
    somente_no_imr = format_out(df_base[~df_base[key].isin(df_ref[key])], cols_ignorar_hora)
    painel_match   = ordenar_por_pedido(painel_match, 'Nº Pedido')
    somente_no_db  = ordenar_por_pedido(somente_no_db, key)
    somente_no_imr = ordenar_por_pedido(somente_no_imr, key)
    df_base_out    = ordenar_por_pedido(format_out(df_base, cols_ignorar_hora), key)
    df_ref_out     = ordenar_por_pedido(format_out(df_ref, cols_ignorar_hora), key)

    total_pedidos_join = len(joined)
    total_divergencias = len(painel_match) if not painel_match.empty else 0
    pct_ok = ((total_pedidos_join - total_divergencias) / total_pedidos_join * 100) if total_pedidos_join > 0 else 0

    return {
        "modo": "pedidos",
        "stats_final": stats_final,
        "painel_match": painel_match,
        "somente_no_db": somente_no_db,
        "somente_no_imr": somente_no_imr,
        "df_base_out": df_base_out,
        "df_ref_out": df_ref_out,
        "comparacao": comparacao,
        "mapa_colunas_imr": mapa_colunas_imr,
        "qtd_linhas_imr": qtd_linhas_imr,
        "qtd_linhas_db": qtd_linhas_db,
        "total_pedidos_join": total_pedidos_join,
        "total_divergencias": total_divergencias,
        "pct_ok": pct_ok,
    }


# =======================================================
# ==== Processamento: Armazém - Quebra Depósito ====
# ==== (chave composta produto + lote, comparação numérica)
# =======================================================
def processar_quebra(df_base, df_ref):
    qtd_linhas_imr = len(df_base)
    qtd_linhas_db  = len(df_ref)

    # Padronização explícita de cabeçalhos
    df_base.columns = [re.sub(r'\s+', ' ', str(c)).strip() for c in df_base.columns]
    mapa_exato = {
        "Nº Produto": "numero_produto",
        "Lote": "codigo_lote",
        "Quantidade Movimentações": "quantidade_movimentacoes",
        "Quantidade Avarias": "quantidade_avarias"
    }
    df_base = df_base.rename(columns=mapa_exato)
    df_base = df_base.loc[:, ~df_base.columns.duplicated()].copy()
    df_ref  = df_ref.loc[:, ~df_ref.columns.duplicated()].copy()

    join_keys = ["numero_produto", "codigo_lote"]
    for k in join_keys:
        if k not in df_base.columns or k not in df_ref.columns:
            raise ValueError(f"Coluna chave '{k}' não encontrada. Verifique os cabeçalhos.")
        df_base[k] = limpar_chave(df_base[k], forte=True)
        df_ref[k]  = limpar_chave(df_ref[k],  forte=True)

    cols_to_compare = ["quantidade_movimentacoes", "quantidade_avarias"]
    for col in cols_to_compare:
        if col in df_base.columns:
            df_base[col] = pd.to_numeric(df_base[col], errors='coerce').fillna(0).round(4)
        if col in df_ref.columns:
            df_ref[col] = pd.to_numeric(df_ref[col], errors='coerce').fillna(0).round(4)

    joined = pd.merge(df_base, df_ref, on=join_keys, how='left', suffixes=('_imr', '_db'))

    def compare_num(col):
        nm_imr, nm_db = f"{col}_imr", f"{col}_db"
        val_imr = joined[nm_imr] if nm_imr in joined.columns else pd.Series([0] * len(joined))
        val_db  = joined[nm_db] if nm_db in joined.columns else pd.Series([np.nan] * len(joined))
        match_final = val_db.notna() & (val_imr == val_db.fillna(0))
        diff = val_db.fillna(0) - val_imr
        df_result = joined[join_keys].copy()
        df_result['coluna'] = col
        df_result['valor_imr'] = val_imr
        df_result['valor_db'] = val_db.fillna(0)
        df_result['diff'] = diff
        df_result['match_final'] = match_final
        return df_result

    comparacao_list = [compare_num(col) for col in cols_to_compare if col in df_base.columns]
    if comparacao_list:
        comparacao = pd.concat(comparacao_list, ignore_index=True)
    else:
        comparacao = pd.DataFrame(columns=join_keys + ['coluna', 'valor_imr', 'valor_db', 'diff', 'match_final'])

    if not comparacao.empty:
        painel_match = comparacao.pivot_table(
            index=join_keys, columns='coluna', values='match_final',
            aggfunc='first', fill_value=False
        ).reset_index()
        detalhe_divergencias = comparacao[comparacao['match_final'] == False]
        colunas_teste = [c for c in painel_match.columns if c not in join_keys]
        mascara_divergencias = (painel_match[colunas_teste] == False).any(axis=1)
        painel_match = painel_match[mascara_divergencias]
        ordem_desejada = join_keys + [c for c in df_base.columns if c in painel_match.columns and c not in join_keys]
        painel_match = painel_match[[c for c in ordem_desejada if c in painel_match.columns]]

        # Normaliza tipos para evitar erro de conversão (Arrow) na exibição do Streamlit:
        # colunas-chave viram texto puro, colunas de divergência viram booleano puro.
        painel_match.columns.name = None
        for c in painel_match.columns:
            if c in join_keys:
                painel_match[c] = painel_match[c].astype(str)
            else:
                painel_match[c] = painel_match[c].astype(bool)
    else:
        painel_match = pd.DataFrame()
        detalhe_divergencias = pd.DataFrame()

    stats_colunas = comparacao.groupby('coluna').agg(
        total_itens_comparados=('match_final', 'count'),
        matches=('match_final', lambda x: x.sum()),
        divergencias=('match_final', lambda x: (~x).sum())
    ).reset_index()
    stats_colunas['pct_assertividade'] = (stats_colunas['matches'] / stats_colunas['total_itens_comparados']) * 100
    stats_colunas['pct_assertividade'] = stats_colunas['pct_assertividade'].apply(lambda x: f"{x:,.2f}%".replace('.', ','))
    stats_colunas['linhas_originais_IMR'] = qtd_linhas_imr
    stats_colunas['linhas_originais_DB']  = qtd_linhas_db

    somente_no_db = df_ref.merge(df_base[join_keys], on=join_keys, how='left', indicator=True) \
                          .query('_merge == "left_only"').drop('_merge', axis=1)
    somente_no_imr = df_base.merge(df_ref[join_keys], on=join_keys, how='left', indicator=True) \
                            .query('_merge == "left_only"').drop('_merge', axis=1)

    def remove_timezone(df):
        df_copy = df.copy()
        for col in df_copy.select_dtypes(include=['datetimetz']).columns:
            df_copy[col] = df_copy[col].dt.tz_localize(None)
        return df_copy

    df_base_out = remove_timezone(df_base)
    df_ref_out  = remove_timezone(df_ref)
    somente_no_db  = remove_timezone(somente_no_db)
    somente_no_imr = remove_timezone(somente_no_imr)

    painel_match         = ordenar_por_chaves(painel_match, join_keys)
    detalhe_divergencias = ordenar_por_chaves(detalhe_divergencias, join_keys)
    somente_no_db        = ordenar_por_chaves(somente_no_db, join_keys)
    somente_no_imr       = ordenar_por_chaves(somente_no_imr, join_keys)
    df_base_out          = ordenar_por_chaves(df_base_out, join_keys)
    df_ref_out           = ordenar_por_chaves(df_ref_out, join_keys)

    # Normaliza tipos de colunas numéricas/booleanas para evitar erro de conversão
    # (Arrow) na exibição do Streamlit.
    if not detalhe_divergencias.empty:
        for c in join_keys:
            if c in detalhe_divergencias.columns:
                detalhe_divergencias[c] = detalhe_divergencias[c].astype(str)
        for c in ['valor_imr', 'valor_db', 'diff']:
            if c in detalhe_divergencias.columns:
                detalhe_divergencias[c] = pd.to_numeric(detalhe_divergencias[c], errors='coerce')
        if 'match_final' in detalhe_divergencias.columns:
            detalhe_divergencias['match_final'] = detalhe_divergencias['match_final'].astype(bool)

    for df_aux in (somente_no_db, somente_no_imr, df_base_out, df_ref_out):
        for c in join_keys:
            if c in df_aux.columns:
                df_aux[c] = df_aux[c].astype(str)

    total_comparacoes = len(comparacao)
    total_matches = int(comparacao['match_final'].sum()) if not comparacao.empty else 0
    pct_ok = (total_matches / total_comparacoes * 100) if total_comparacoes > 0 else 0
    total_divergencias = len(painel_match) if not painel_match.empty else 0

    return {
        "modo": "quebra",
        "stats_final": stats_colunas,
        "painel_match": painel_match,
        "detalhe_divergencias": detalhe_divergencias,
        "somente_no_db": somente_no_db,
        "somente_no_imr": somente_no_imr,
        "df_base_out": df_base_out,
        "df_ref_out": df_ref_out,
        "comparacao": comparacao,
        "mapa_colunas_imr": {},
        "qtd_linhas_imr": qtd_linhas_imr,
        "qtd_linhas_db": qtd_linhas_db,
        "total_pedidos_join": len(joined),
        "total_divergencias": total_divergencias,
        "pct_ok": pct_ok,
    }


# =======================================================
# ==== Geração do Excel (genérica: dict de abas) ====
# =======================================================
def gerar_excel(sheets):
    output = io.BytesIO()
    header_fill = PatternFill(start_color="002060", end_color="002060", fill_type="solid")
    header_font = XLFont(color="FFFFFF", bold=True)

    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        for sheet_name, df in sheets.items():
            df.to_excel(writer, sheet_name=sheet_name, index=False)

        for sheet_name in writer.sheets:
            ws = writer.sheets[sheet_name]
            for cell in ws[1]:
                cell.fill = header_fill
                cell.font = header_font
            for col in ws.columns:
                max_length = 0
                for cell in col:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                ws.column_dimensions[col[0].column_letter].width = max_length + 2

    output.seek(0)
    return output


# =======================================================
# ==== Helpers Word ====
# =======================================================
def set_cell_bg(cell, rgb_tuple):
    r, g, b = rgb_tuple
    hex_color = f"{r:02X}{g:02X}{b:02X}"
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(*COR_HEADER_AZUL)
    return p


def add_table_word(doc, df, max_rows=200, col_width_twips=None):
    """Adiciona tabela ocupando toda a largura da página."""
    if df.empty:
        doc.add_paragraph("Nenhum dado encontrado.")
        return

    df_show = df.head(max_rows).fillna('').astype(str)
    cols = list(df_show.columns)
    n_cols = len(cols)

    PAGE_W = 9072
    w = col_width_twips if col_width_twips else PAGE_W // n_cols

    table = doc.add_table(rows=1 + len(df_show), cols=n_cols)
    table.style = 'Table Grid'

    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(PAGE_W))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    for i, col_name in enumerate(cols):
        cell = table.rows[0].cells[i]
        cell.text = col_name
        cell.width = w
        set_cell_bg(cell, COR_HEADER_AZUL)
        run = cell.paragraphs[0].runs[0]
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True
        run.font.size = Pt(8)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx, (_, row) in enumerate(df_show.iterrows()):
        bg = COR_CINZA_CLARO if row_idx % 2 == 0 else (255, 255, 255)
        for col_idx, val in enumerate(row):
            cell = table.rows[row_idx + 1].cells[col_idx]
            val_str = str(val)
            cell.text = val_str
            cell.width = w
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(val_str)
            run.font.size = Pt(8)

    if len(df) > max_rows:
        doc.add_paragraph(f"⚠ Exibindo {max_rows} de {len(df)} linhas. Consulte o Excel para a tabela completa.")


# =======================================================
# ==== Geração do Word (relatórios por Nº Pedido) ====
# =======================================================
def gerar_word(tipo_relatorio, stats_final, painel_match, somente_no_db, somente_no_imr,
               qtd_linhas_imr, qtd_linhas_db, total_pedidos_join,
               total_divergencias, pct_ok, comparacao):
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)

    hoje_str = datetime.datetime.now().strftime("%d/%m/%Y %H:%M")

    titulo = doc.add_heading('Relatório de Cruzamento De-Para', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.color.rgb = RGBColor(*COR_HEADER_AZUL)
        run.font.size = Pt(18)
    sub = doc.add_paragraph(f'{tipo_relatorio}  |  IMR vs Databricks  |  Gerado em {hoje_str}')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(89, 89, 89)
    sub.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    add_heading(doc, '1. Resumo Executivo', level=1)
    resumo_data = [
        ("Total de linhas no IMR",       str(qtd_linhas_imr)),
        ("Total de linhas no Databricks", str(qtd_linhas_db)),
        ("Pedidos em comum (Inner Join)", str(total_pedidos_join)),
        ("Pedidos com divergência",       str(total_divergencias)),
        ("Pedidos somente no DB",         str(len(somente_no_db))),
        ("Pedidos somente no IMR",        str(len(somente_no_imr))),
        ("% de Conformidade",             f"{pct_ok:.1f}%"),
    ]
    t = doc.add_table(rows=len(resumo_data), cols=2)
    t.style = 'Table Grid'
    for i, (label, valor) in enumerate(resumo_data):
        bg = COR_CINZA_CLARO if i % 2 == 0 else (255, 255, 255)
        t.rows[i].cells[0].text = label
        set_cell_bg(t.rows[i].cells[0], bg)
        t.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        t.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        t.rows[i].cells[0].width = Inches(4)
        t.rows[i].cells[1].text = valor
        if label == "% de Conformidade":
            val_cor = COR_OK if pct_ok >= 95 else (COR_AVISO if pct_ok >= 80 else COR_DIVERGENCIA)
            set_cell_bg(t.rows[i].cells[1], val_cor)
        elif label == "Pedidos com divergência" and total_divergencias > 0:
            set_cell_bg(t.rows[i].cells[1], COR_DIVERGENCIA)
        else:
            set_cell_bg(t.rows[i].cells[1], bg)
        t.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        t.rows[i].cells[1].width = Inches(2)

    doc.add_paragraph()
    if pct_ok >= 95:
        msg = f"✅ Conformidade em {pct_ok:.1f}%, dentro do padrão esperado."
    elif pct_ok >= 80:
        msg = f"⚠ Conformidade em {pct_ok:.1f}%. Recomenda-se revisão das divergências."
    else:
        msg = f"❌ Conformidade em {pct_ok:.1f}%. Ação corretiva necessária."
    p = doc.add_paragraph(msg)
    p.runs[0].font.italic = True
    p.runs[0].font.size = Pt(10)
    doc.add_page_break()

    add_heading(doc, '2. Estatísticas por Coluna', level=1)
    doc.add_paragraph('Total de comparações por campo, com percentual de conformidade e divergência.')
    doc.add_paragraph()
    stats_display = stats_final.copy()
    for col in ['total_comparacoes', 'verdadeiro_n', 'falso_n']:
        if col in stats_display.columns:
            stats_display[col] = stats_display[col].apply(
                lambda x: str(int(x)) if pd.notna(x) else ''
            )
    add_table_word(doc, stats_display)
    doc.add_page_break()

    add_heading(doc, '3. Divergências Detalhadas por Campo', level=1)
    doc.add_paragraph('Pedidos onde os valores do IMR e do Databricks não coincidem, agrupados por campo.')
    if not comparacao.empty:
        colunas_div = comparacao[~comparacao['match_final']]['coluna'].unique()
        for col_div in sorted(colunas_div):
            df_col = comparacao[
                (comparacao['coluna'] == col_div) & (~comparacao['match_final'])
            ][['Nº Pedido', 'valor_imr', 'valor_db']].copy()
            df_col.columns = ['Nº Pedido', 'Valor IMR', 'Valor Databricks']
            add_heading(doc, f'Campo: {col_div}  ({len(df_col)} divergência(s))', level=2)
            add_table_word(doc, df_col, max_rows=100)
            doc.add_paragraph()
    else:
        doc.add_paragraph("Nenhuma divergência encontrada.")
    doc.add_page_break()

    add_heading(doc, '4. Pedidos Somente no Databricks', level=1)
    doc.add_paragraph('Pedidos no Databricks sem correspondência no IMR.')
    doc.add_paragraph()
    key = 'numero_pedido'
    db_simples = somente_no_db[[key]].rename(columns={key: 'Nº Pedido'}) if key in somente_no_db.columns else somente_no_db[['Nº Pedido']] if 'Nº Pedido' in somente_no_db.columns else somente_no_db.iloc[:, :1]
    add_table_word(doc, db_simples, max_rows=200)
    doc.add_page_break()

    add_heading(doc, '5. Pedidos Somente no IMR', level=1)
    doc.add_paragraph('Pedidos no IMR sem correspondência no Databricks.')
    doc.add_paragraph()
    imr_simples = somente_no_imr[[key]].rename(columns={key: 'Nº Pedido'}) if key in somente_no_imr.columns else somente_no_imr[['Nº Pedido']] if 'Nº Pedido' in somente_no_imr.columns else somente_no_imr.iloc[:, :1]
    add_table_word(doc, imr_simples, max_rows=200)
    doc.add_page_break()

    add_heading(doc, '6. Painel de Match por Pedido', level=1)
    doc.add_paragraph('Pedidos com ao menos uma divergência e quais campos estão em conformidade.')
    doc.add_paragraph()
    if painel_match.empty:
        doc.add_paragraph("Nenhum pedido com divergência.")
    else:
        pm = painel_match.copy()
        for c in pm.columns:
            if pm[c].dtype == bool:
                pm[c] = pm[c].map({True: '✔', False: '✘'})
        add_table_word(doc, pm, max_rows=150)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# =======================================================
# ==== Geração do Word (Armazém - Quebra Depósito) ====
# =======================================================
def gerar_word_quebra(stats_final, painel_match, detalhe_divergencias,
                      somente_no_db, somente_no_imr,
                      qtd_linhas_imr, qtd_linhas_db, total_itens,
                      total_divergencias, pct_ok):
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)

    hoje_str = datetime.datetime.now().strftime("%d/%m/%Y %H:%M")

    titulo = doc.add_heading('Relatório de Cruzamento De-Para', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.color.rgb = RGBColor(*COR_HEADER_AZUL)
        run.font.size = Pt(18)
    sub = doc.add_paragraph(f'Armazém - Quebra Depósito  |  IMR vs Databricks  |  Gerado em {hoje_str}')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(89, 89, 89)
    sub.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    add_heading(doc, '1. Resumo Executivo', level=1)
    resumo_data = [
        ("Total de linhas no IMR",             str(qtd_linhas_imr)),
        ("Total de linhas no Databricks",      str(qtd_linhas_db)),
        ("Itens cruzados (Produto + Lote)",    str(total_itens)),
        ("Itens com divergência",              str(total_divergencias)),
        ("Itens somente no DB",                str(len(somente_no_db))),
        ("Itens somente no IMR",               str(len(somente_no_imr))),
        ("% de Assertividade",                 f"{pct_ok:.1f}%".replace('.', ',')),
    ]
    t = doc.add_table(rows=len(resumo_data), cols=2)
    t.style = 'Table Grid'
    for i, (label, valor) in enumerate(resumo_data):
        bg = COR_CINZA_CLARO if i % 2 == 0 else (255, 255, 255)
        t.rows[i].cells[0].text = label
        set_cell_bg(t.rows[i].cells[0], bg)
        t.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        t.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        t.rows[i].cells[0].width = Inches(4)
        t.rows[i].cells[1].text = valor
        if label == "% de Assertividade":
            val_cor = COR_OK if pct_ok >= 95 else (COR_AVISO if pct_ok >= 80 else COR_DIVERGENCIA)
            set_cell_bg(t.rows[i].cells[1], val_cor)
        elif label == "Itens com divergência" and total_divergencias > 0:
            set_cell_bg(t.rows[i].cells[1], COR_DIVERGENCIA)
        else:
            set_cell_bg(t.rows[i].cells[1], bg)
        t.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        t.rows[i].cells[1].width = Inches(2)

    doc.add_paragraph()
    pct_str = f"{pct_ok:.1f}".replace('.', ',')
    if pct_ok >= 95:
        msg = f"✅ Assertividade em {pct_str}%, dentro do padrão esperado."
    elif pct_ok >= 80:
        msg = f"⚠ Assertividade em {pct_str}%. Recomenda-se revisão das divergências."
    else:
        msg = f"❌ Assertividade em {pct_str}%. Ação corretiva necessária."
    p = doc.add_paragraph(msg)
    p.runs[0].font.italic = True
    p.runs[0].font.size = Pt(10)
    doc.add_page_break()

    add_heading(doc, '2. Estatísticas por Coluna', level=1)
    doc.add_paragraph('Total de itens comparados por campo numérico, com percentual de assertividade.')
    doc.add_paragraph()
    stats_display = stats_final.copy()
    for col in ['total_itens_comparados', 'matches', 'divergencias',
                'linhas_originais_IMR', 'linhas_originais_DB']:
        if col in stats_display.columns:
            stats_display[col] = stats_display[col].apply(
                lambda x: str(int(x)) if pd.notna(x) else ''
            )
    add_table_word(doc, stats_display)
    doc.add_page_break()

    add_heading(doc, '3. Detalhe das Divergências', level=1)
    doc.add_paragraph(
        'Itens (Produto + Lote) onde as quantidades do IMR e do Databricks não coincidem. '
        'A coluna "diff" mostra a diferença (DB - IMR).'
    )
    doc.add_paragraph()
    if detalhe_divergencias.empty:
        doc.add_paragraph("Nenhuma divergência encontrada.")
    else:
        det = detalhe_divergencias.copy()
        cols_renomear = {
            'numero_produto': 'Nº Produto', 'codigo_lote': 'Lote',
            'coluna': 'Campo', 'valor_imr': 'Valor IMR',
            'valor_db': 'Valor DB', 'diff': 'Diferença', 'match_final': 'Match'
        }
        det = det.rename(columns={k: v for k, v in cols_renomear.items() if k in det.columns})
        if 'Match' in det.columns:
            det = det.drop(columns=['Match'])
        add_table_word(doc, det, max_rows=150)
    doc.add_page_break()

    add_heading(doc, '4. Itens Somente no Databricks', level=1)
    doc.add_paragraph('Itens (Produto + Lote) no Databricks sem correspondência no IMR.')
    doc.add_paragraph()
    keys_quebra = ['numero_produto', 'codigo_lote']
    db_cols = [k for k in keys_quebra if k in somente_no_db.columns]
    db_simples = somente_no_db[db_cols].rename(
        columns={'numero_produto': 'Nº Produto', 'codigo_lote': 'Lote'}
    ) if db_cols else somente_no_db.iloc[:, :2]
    add_table_word(doc, db_simples, max_rows=200)
    doc.add_page_break()

    add_heading(doc, '5. Itens Somente no IMR', level=1)
    doc.add_paragraph('Itens (Produto + Lote) no IMR sem correspondência no Databricks.')
    doc.add_paragraph()
    imr_cols = [k for k in keys_quebra if k in somente_no_imr.columns]
    imr_simples = somente_no_imr[imr_cols].rename(
        columns={'numero_produto': 'Nº Produto', 'codigo_lote': 'Lote'}
    ) if imr_cols else somente_no_imr.iloc[:, :2]
    add_table_word(doc, imr_simples, max_rows=200)
    doc.add_page_break()

    add_heading(doc, '6. Painel de Match por Item', level=1)
    doc.add_paragraph('Itens com ao menos uma divergência e quais campos estão em conformidade.')
    doc.add_paragraph()
    if painel_match.empty:
        doc.add_paragraph("Nenhum item com divergência.")
    else:
        pm = painel_match.copy()
        for c in pm.columns:
            if pm[c].dtype == bool:
                pm[c] = pm[c].map({True: '✔', False: '✘'})
        add_table_word(doc, pm, max_rows=150)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# =======================================================
# ==== Interface ====
# =======================================================
if not TOKEN:
    st.error("⚠️ Token do Databricks não encontrado. Crie o arquivo `.env` com DATABRICKS_TOKEN=seu_token.")
    st.stop()

st.divider()
st.subheader("1. Selecione o tipo de relatório")

tipo_relatorio = st.selectbox(
    "Tipo de relatório",
    options=list(CONFIG.keys()),
    format_func=lambda t: f"{CONFIG[t]['emoji']} {t}"
)
cfg = CONFIG[tipo_relatorio]

st.divider()
st.subheader("2. Selecione o período de referência")

col1, col2 = st.columns(2)
with col1:
    ano_input = st.selectbox("Ano", options=list(range(2025, 2029)), index=0)
with col2:
    mes_input = st.selectbox(
        "Mês",
        options=list(range(1, 13)),
        format_func=lambda m: NOMES_MESES[m - 1]
    )

nome_mes = NOMES_MESES[mes_input - 1]

# Validação específica do Comprovantes (dicionário de referências IMR)
periodo_valido = True
if tipo_relatorio == "Comprovantes":
    chave_busca = f"{ano_input}-{mes_input:02d}"
    if chave_busca in DICIONARIO_IMR_COMPROVANTES:
        imr_ref = DICIONARIO_IMR_COMPROVANTES[chave_busca]
        st.success(f"Período selecionado: **{nome_mes} de {ano_input}** (Referência IMR: {imr_ref})")
    else:
        periodo_valido = False
        st.error(
            f"O período {nome_mes} de {ano_input} ainda não está cadastrado no dicionário de "
            f"referências IMR de Comprovantes. Fale com o responsável para atualizar o app."
        )
else:
    st.success(f"Período selecionado: **{nome_mes} de {ano_input}**")

st.divider()
st.subheader("3. Upload do arquivo IMR")

with st.expander("📋 Instruções: como preparar o arquivo antes do upload", expanded=False):
    st.markdown(f"""
O relatório bruto do IMR vem com abas mescladas, logo, cabeçalho e colunas vazias.
Siga os passos abaixo **na aba "{cfg['aba_bruta']}"** do arquivo bruto antes de subir aqui.
""")

    st.markdown("**Passo 1 — Desagrupar as abas**")
    st.markdown(
        "Ao abrir o arquivo, várias abas podem vir selecionadas juntas (em branco). "
        f"Segure **CTRL** e clique nas abas para desagrupá-las, deixando apenas a aba "
        f"**{cfg['aba_bruta']}** ativa."
    )
    col_a, col_b = st.columns(2)
    with col_a:
        st.image("assets_manual/passo_01.png", caption="Abas agrupadas (errado)")
    with col_b:
        st.image("assets_manual/passo_02.png", caption="Apenas uma aba ativa (correto)")

    st.markdown("**Passo 2 — Remover mesclagem de células**")
    st.markdown(
        "Clique no triângulo no canto superior esquerdo da planilha (entre a coluna A "
        "e a linha 1) para selecionar tudo. Na guia **Página Inicial**, clique em "
        "**Mesclar e Centralizar** para desativar a mesclagem. Pode levar alguns "
        "segundos para processar."
    )
    col_c, col_d = st.columns(2)
    with col_c:
        st.image("assets_manual/passo_03.png", caption="Selecionar a planilha inteira")
    with col_d:
        st.image("assets_manual/passo_04.png", caption="Desativar mesclagem")

    st.markdown("**Passo 3 — Limpar o layout**")
    st.markdown(
        "- Arraste a divisória entre as colunas A e B para expandir e visualizar melhor os dados.\n"
        "- Clique com o botão direito sobre o logo da VTCLOG e selecione **Recortar** (ou delete).\n"
        "- Selecione as linhas de cabeçalho (geralmente 1 a 5), clique com o botão direito "
        "sobre os números das linhas e escolha **Excluir**."
    )
    col_e, col_f, col_g = st.columns(3)
    with col_e:
        st.image("assets_manual/passo_05.png", caption="Ajustar colunas")
    with col_f:
        st.image("assets_manual/passo_06.png", caption="Remover logo")
    with col_g:
        st.image("assets_manual/passo_07.png", caption="Remover cabeçalho")

    st.markdown("**Passo 4 — Remover colunas vazias**")
    st.markdown(
        "Segure **CTRL** e selecione todas as colunas visivelmente em branco. Clique com "
        "o botão direito no cabeçalho de uma delas e escolha **Excluir**."
    )
    st.image("assets_manual/passo_08.png", caption="Selecionar e excluir colunas vazias")

    st.markdown("**Passo 5 — Filtrar os dados**")
    st.markdown(
        "Selecione a planilha inteira novamente, vá em **Dados > Filtro**. No filtro da "
        "primeira coluna, desmarque as opções em branco (e o nome do cabeçalho repetido) "
        "para manter apenas as linhas com dados."
    )
    st.image("assets_manual/passo_09.png", caption="Filtrar linhas vazias")

    st.markdown("**Passo 6 — Exportar o arquivo limpo**")
    st.markdown(
        "Copie os dados filtrados (**CTRL+C**), cole em um Excel novo (**CTRL+V**) e "
        f"salve com o nome padrão (ex: `{cfg['exemplo_arquivo']}`). "
        "Esse é o arquivo que você deve subir no campo abaixo."
    )
    col_h, col_i = st.columns(2)
    with col_h:
        st.image("assets_manual/passo_10.png", caption="Copiar dados filtrados")
    with col_i:
        st.image("assets_manual/passo_11.png", caption="Salvar arquivo final")

    st.info("💡 Dúvidas ou dificuldades? Fale com o responsável pelo processo antes de prosseguir.")

file_imr = st.file_uploader(
    f"**📄 Base IMR — {tipo_relatorio}** (Ex: {cfg['exemplo_arquivo']})",
    type=["xlsx", "xls"]
)

st.divider()

# Inicializa session_state
if "resultado" not in st.session_state:
    st.session_state.resultado = None

if file_imr and periodo_valido:
    col_btn1, col_btn2 = st.columns([3, 1])
    with col_btn1:
        processar = st.button("🔄 Buscar dados do Databricks e processar", type="primary", use_container_width=True)
    with col_btn2:
        if st.button("🗑️ Resetar", use_container_width=True):
            st.session_state.resultado = None
            st.rerun()

    if processar:
        with st.spinner(f"Conectando ao Databricks ({tipo_relatorio} — {nome_mes} de {ano_input})..."):
            try:
                query = montar_query(tipo_relatorio, ano_input, mes_input)
                df_ref = extrair_databricks(query)
                st.success(f"✅ {f'{len(df_ref):,}'.replace(',', '.')} linhas extraídas do Databricks.")
            except Exception as e:
                st.error(f"Erro ao conectar ao Databricks: {e}")
                st.stop()

        with st.spinner("Processando os arquivos..."):
            try:
                df_base = pd.read_excel(file_imr)
                if tipo_relatorio == "Armazém - Quebra Depósito":
                    resultado = processar_quebra(df_base, df_ref)
                else:
                    resultado = processar_pedidos(df_base, df_ref, cfg)
            except ValueError as e:
                st.error(str(e))
                st.stop()

            resultado["tipo"] = tipo_relatorio
            resultado["ano_input"] = ano_input
            resultado["mes_input"] = mes_input
            st.session_state.resultado = resultado

# Exibe resultados se existirem e forem do tipo atualmente selecionado
r = st.session_state.resultado
if r and r.get("tipo") != tipo_relatorio:
    st.info(
        f"ℹ️ Há um resultado processado de **{r.get('tipo')}** nesta sessão. "
        "Selecione esse tipo novamente para vê-lo, ou processe um novo arquivo."
    )
    r = None

if r:
    st.success(f"Processamento concluído! ({r['tipo']} — {NOMES_MESES[r['mes_input'] - 1]} de {r['ano_input']})")

    if r["modo"] == "pedidos":
        m1, m2, m3, m4, m5 = st.columns(5)
        m1.metric("Linhas IMR",         f"{r['qtd_linhas_imr']:,}".replace(',', '.'))
        m2.metric("Linhas DB",          f"{r['qtd_linhas_db']:,}".replace(',', '.'))
        m3.metric("Match (Inner Join)", f"{r['total_pedidos_join']:,}".replace(',', '.'))
        m4.metric("Com Divergência",    f"{r['total_divergencias']:,}".replace(',', '.'))
        m5.metric("% Conformidade",     f"{r['pct_ok']:.1f}%".replace('.', ','))
    else:
        m1, m2, m3, m4, m5 = st.columns(5)
        m1.metric("Linhas IMR",           f"{r['qtd_linhas_imr']:,}".replace(',', '.'))
        m2.metric("Linhas DB",            f"{r['qtd_linhas_db']:,}".replace(',', '.'))
        m3.metric("Itens comparados",     f"{r['total_pedidos_join']:,}".replace(',', '.'))
        m4.metric("Com Divergência",      f"{r['total_divergencias']:,}".replace(',', '.'))
        m5.metric("% Assertividade",      f"{r['pct_ok']:.1f}%".replace('.', ','))

    st.divider()

    if r["modo"] == "pedidos":
        tab1, tab2, tab3, tab4 = st.tabs([
            "📊 Stats por Coluna",
            "⚠️ Painel Match (Divergências)",
            "🔵 Somente no DB",
            "🟡 Somente no IMR"
        ])
        with tab1:
            st.dataframe(r['stats_final'], use_container_width=True, hide_index=True)
        with tab2:
            if not r['painel_match'].empty:
                st.dataframe(r['painel_match'], use_container_width=True, hide_index=True)
            else:
                st.info("Nenhuma divergência.")
        with tab3:
            if not r['somente_no_db'].empty:
                st.dataframe(r['somente_no_db'], use_container_width=True, hide_index=True)
            else:
                st.info("Nenhum pedido exclusivo no DB.")
        with tab4:
            if not r['somente_no_imr'].empty:
                st.dataframe(r['somente_no_imr'], use_container_width=True, hide_index=True)
            else:
                st.info("Nenhum pedido exclusivo no IMR.")
    else:
        tab1, tab2, tab3, tab4, tab5 = st.tabs([
            "📊 Estatísticas",
            "⚠️ Painel Match (Divergências)",
            "🔍 Detalhe Divergências",
            "🔵 Somente no DB",
            "🟡 Somente no IMR"
        ])
        with tab1:
            st.dataframe(r['stats_final'], use_container_width=True, hide_index=True)
        with tab2:
            if not r['painel_match'].empty:
                st.dataframe(r['painel_match'], use_container_width=True, hide_index=True)
            else:
                st.info("Nenhuma divergência.")
        with tab3:
            if not r['detalhe_divergencias'].empty:
                st.dataframe(r['detalhe_divergencias'], use_container_width=True, hide_index=True)
            else:
                st.info("Nenhuma divergência detalhada.")
        with tab4:
            if not r['somente_no_db'].empty:
                st.dataframe(r['somente_no_db'], use_container_width=True, hide_index=True)
            else:
                st.info("Nenhum item exclusivo no DB.")
        with tab5:
            if not r['somente_no_imr'].empty:
                st.dataframe(r['somente_no_imr'], use_container_width=True, hide_index=True)
            else:
                st.info("Nenhum item exclusivo no IMR.")

    if r['mapa_colunas_imr']:
        with st.expander("🔍 Colunas identificadas automaticamente (Fuzzy Match)"):
            for original, nova in r['mapa_colunas_imr'].items():
                st.write(f"• `{original}` → `{nova}`")

    st.divider()
    st.subheader("📥 Downloads")
    hoje = datetime.datetime.now().strftime("%Y-%m-%d")
    prefixo = CONFIG[r['tipo']]['prefixo_saida']
    nome_base = f"({hoje}) {prefixo}_{r['ano_input']}_{r['mes_input']:02d}"

    if r["modo"] == "pedidos":
        col_dl1, col_dl2 = st.columns(2)
        with col_dl1:
            excel_bytes = gerar_excel({
                'Painel_Match':       r['painel_match'],
                'Stats_%_por_Coluna': r['stats_final'],
                'Somente_no_DB':      r['somente_no_db'],
                'Somente_no_IMR':     r['somente_no_imr'],
                'Base_IMR_Limpa':     r['df_base_out'],
                'Base_DB':            r['df_ref_out'],
            })
            st.download_button(
                label="⬇️ Baixar relatório Excel",
                data=excel_bytes,
                file_name=f"{nome_base}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
        with col_dl2:
            word_bytes = gerar_word(
                r['tipo'],
                r['stats_final'], r['painel_match'], r['somente_no_db'], r['somente_no_imr'],
                r['qtd_linhas_imr'], r['qtd_linhas_db'], r['total_pedidos_join'],
                r['total_divergencias'], r['pct_ok'], r['comparacao']
            )
            st.download_button(
                label="⬇️ Baixar relatório Word",
                data=word_bytes,
                file_name=f"({hoje}) Relatorio_{prefixo}_{r['ano_input']}_{r['mes_input']:02d}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary"
            )
    else:
        col_dl1, col_dl2 = st.columns(2)
        with col_dl1:
            excel_bytes = gerar_excel({
                'Resumo_Match':          r['painel_match'],
                'Detalhe_Divergencias':  r['detalhe_divergencias'],
                'Estatisticas':          r['stats_final'],
                'somente_no_databricks': r['somente_no_db'],
                'somente_no_imr':        r['somente_no_imr'],
                'Base_Bruta_IMR':        r['df_base_out'],
                'Base_Bruta_DB':         r['df_ref_out'],
            })
            st.download_button(
                label="⬇️ Baixar relatório Excel",
                data=excel_bytes,
                file_name=f"{nome_base}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
        with col_dl2:
            word_bytes = gerar_word_quebra(
                r['stats_final'], r['painel_match'], r['detalhe_divergencias'],
                r['somente_no_db'], r['somente_no_imr'],
                r['qtd_linhas_imr'], r['qtd_linhas_db'], r['total_pedidos_join'],
                r['total_divergencias'], r['pct_ok']
            )
            st.download_button(
                label="⬇️ Baixar relatório Word",
                data=word_bytes,
                file_name=f"({hoje}) Relatorio_{prefixo}_{r['ano_input']}_{r['mes_input']:02d}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary"
            )
else:
    if not st.session_state.resultado:
        st.info("⬆️ Faça o upload do arquivo IMR para continuar.")